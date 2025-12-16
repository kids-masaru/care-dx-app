"""
Create Word document for care-dx-app manual with screenshots
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
import os

# Create document
doc = Document()

# Set up styles
style = doc.styles['Title']
style.font.size = Pt(28)
style.font.color.rgb = RGBColor(14, 165, 233)

style = doc.styles['Heading 1']
style.font.size = Pt(18)
style.font.color.rgb = RGBColor(14, 165, 233)

style = doc.styles['Heading 2']
style.font.size = Pt(14)
style.font.color.rgb = RGBColor(30, 41, 59)

# Title
title = doc.add_heading('介護DX カカナイ', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

subtitle = doc.add_paragraph('操作マニュアル v1.0')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()  # Spacer

# Table of Contents
doc.add_heading('目次', level=1)
toc = doc.add_paragraph()
toc.add_run('1. はじめに\n')
toc.add_run('2. 運営会議録の作成\n')
toc.add_run('3. サービス担当者会議議事録の作成\n')
toc.add_run('4. アセスメントシートの作成\n')
toc.add_run('5. よくある質問（FAQ）')

doc.add_page_break()

# Section 1: Introduction
doc.add_heading('1. はじめに', level=1)

doc.add_heading('アプリの概要', level=2)
doc.add_paragraph(
    '「介護DX カカナイ」は、介護業務のドキュメント作成を自動化するWebアプリケーションです。'
    '音声データやPDF書類をアップロードするだけで、AIが自動的に情報を抽出し、必要な書類を作成します。'
)

doc.add_heading('3つの主要機能', level=2)
features = doc.add_paragraph()
features.add_run('🎙️ 運営会議録\n').bold = True
features.add_run('　音声データから議事録を自動作成。Google Driveに自動保存。\n\n')
features.add_run('🎙️ サービス担当者会議議事録\n').bold = True
features.add_run('　音声データから詳細な議事録を作成。各事業所の役割分担も抽出。\n\n')
features.add_run('📄 アセスメントシート\n').bold = True
features.add_run('　PDFから情報を抽出しスプレッドシートに自動転記。')

doc.add_heading('必要な準備', level=2)
prep = doc.add_paragraph()
prep.add_run('• 音声データ：').bold = True
prep.add_run('スマートフォンやICレコーダーで録音したファイル（MP3, M4A, WAV形式）\n')
prep.add_run('• PDF書類：').bold = True
prep.add_run('介護保険証、主治医意見書などのスキャンデータ\n')
prep.add_run('• Googleアカウント：').bold = True
prep.add_run('スプレッドシートへの書き込みに使用')

doc.add_page_break()

# Section 2: Management Meeting
doc.add_heading('2. 運営会議録の作成', level=1)

# Screenshot
img_path = r'C:/Users/700289/.gemini/antigravity/brain/fbf3d6ff-07ff-47b4-8e70-3b8c3ea2779a/uploaded_image_1_1765334026818.png'
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    caption = doc.add_paragraph('【運営会議録の入力画面】')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('操作手順', level=2)

steps = [
    ('1. シートタイプを選択', 'サイドバーの「対象シート（機能選択）」ドロップダウンから「運営会議録」を選択します。'),
    ('2. 会議情報を入力', '開催日、開催場所、参加者、開催時間（開始〜終了）を入力します。'),
    ('3. 音声ファイルをアップロード', '「音声ファイルのアップロード」エリアにファイルをドラッグ＆ドロップ、または「Browse files」ボタンでファイルを選択します。'),
    ('4. AI処理を実行', '「AI処理を実行」ボタンをクリックします。処理完了まで数分かかる場合があります。'),
    ('5. 結果を確認', '抽出された議事録が表示され、自動的にGoogleスプレッドシートに転記されます。'),
]

for title, desc in steps:
    p = doc.add_paragraph()
    p.add_run(title + '\n').bold = True
    p.add_run(desc)

tip = doc.add_paragraph()
tip.add_run('💡 自動保存機能：').bold = True
tip.add_run('アップロードした音声ファイルは、指定のGoogle Driveフォルダに自動保存されます。ファイル名は「YYYYMMDD_HHMMSS_元のファイル名」形式です。')

doc.add_page_break()

# Section 3: Service Meeting
doc.add_heading('3. サービス担当者会議議事録の作成', level=1)

# Screenshot
img_path = r'C:/Users/700289/.gemini/antigravity/brain/fbf3d6ff-07ff-47b4-8e70-3b8c3ea2779a/uploaded_image_2_1765334026818.png'
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    caption = doc.add_paragraph('【サービス担当者会議議事録の入力画面】')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('操作手順', level=2)

steps = [
    ('1. シートタイプを選択', 'サイドバーで「サービス担当者会議議事録」を選択します。'),
    ('2. 会議情報を入力', '担当者名、利用者名、開催場所、開催時間、開催日、開催回数を入力します。'),
    ('3. 音声ファイルをアップロード', '会議を録音した音声ファイルを選択します。'),
    ('4. AI処理を実行', '「AI処理を実行」ボタンをクリックします。'),
]

for title, desc in steps:
    p = doc.add_paragraph()
    p.add_run(title + '\n').bold = True
    p.add_run(desc)

doc.add_heading('出力フォーマット', level=2)
doc.add_paragraph('議事録は以下の形式で作成されます：')

format_text = """①【本人及び家族の意向】
　・本人⇒「自分でできることは自分でやりたい」
　・家族⇒「安全に過ごしてほしい」

②【心身・生活状況】
　・身体状況⇒...
　・精神状況⇒...
　・生活状況⇒...

③【会議の結論・ケアプラン詳細】
　・主な検討事項と結論：...

④【各事業所の役割分担と確認事項】
　＊デイサービスA⇒...
　＊訪問看護B⇒...

⑤【福祉用具・住宅改修等に関する検討事項】
　・現状の課題：...
　・検討内容と経緯：...
　・結論：..."""

format_para = doc.add_paragraph(format_text)
format_para.style = 'Quote'

doc.add_page_break()

# Section 4: Assessment Sheet
doc.add_heading('4. アセスメントシートの作成', level=1)

# Screenshot
img_path = r'C:/Users/700289/.gemini/antigravity/brain/fbf3d6ff-07ff-47b4-8e70-3b8c3ea2779a/uploaded_image_0_1765334026818.png'
if os.path.exists(img_path):
    doc.add_picture(img_path, width=Inches(6))
    caption = doc.add_paragraph('【アセスメントシートの入力画面】')
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_heading('操作手順', level=2)

steps = [
    ('1. シートタイプを選択', 'サイドバーで「アセスメントシート」を選択します。'),
    ('2. PDFファイルをアップロード', '介護保険証、主治医意見書などのPDFをアップロードします。複数ファイルを同時にアップロード可能です。'),
    ('3. スプレッドシートIDを入力', '転記先のGoogleスプレッドシートのIDを入力します。URLの「/d/」と「/edit」の間の文字列がIDです。'),
    ('4. AI処理を実行', '「AI処理を実行」ボタンをクリックします。'),
]

for title, desc in steps:
    p = doc.add_paragraph()
    p.add_run(title + '\n').bold = True
    p.add_run(desc)

doc.add_heading('抽出される情報', level=2)
info_list = doc.add_paragraph()
info_list.add_run('• 利用者基本情報：氏名、生年月日、住所、電話番号\n')
info_list.add_run('• 介護保険情報：被保険者番号、要介護度、認定有効期間\n')
info_list.add_run('• 主治医情報：医療機関名、医師名、連絡先\n')
info_list.add_run('• 既往歴・現病歴：疾病名、治療状況\n')
info_list.add_run('• ADL・IADL情報：移動、食事、入浴、排泄等の自立度')

doc.add_page_break()

# Section 5: FAQ
doc.add_heading('5. よくある質問（FAQ）', level=1)

faqs = [
    ('Q: 音声ファイルの長さに制限はありますか？', 'A: 長時間の音声でも処理可能ですが、1時間を超える場合は処理に時間がかかります。'),
    ('Q: 対応している音声形式は？', 'A: MP3, M4A, WAV形式に対応しています。'),
    ('Q: PDFの文字が読み取れない場合は？', 'A: スキャン品質が低いとOCR精度が下がります。できるだけ高解像度（300dpi以上）でスキャンしてください。'),
    ('Q: 抽出結果を修正できますか？', 'A: 転記先のGoogleスプレッドシートで直接編集できます。'),
    ('Q: エラーが発生した場合は？', 'A: ページをリロードして再度お試しください。それでも解決しない場合は管理者にお問い合わせください。'),
    ('Q: アップロードしたファイルはどこに保存されますか？', 'A: 運営会議・サービス担当者会議の音声ファイルは、設定で指定したGoogle Driveフォルダに自動保存されます。'),
]

for q, a in faqs:
    p = doc.add_paragraph()
    p.add_run(q + '\n').bold = True
    p.add_run(a)
    doc.add_paragraph()  # Spacer

# Footer
doc.add_paragraph()
footer = doc.add_paragraph('介護DX カカナイ v1.0 | Powered by Google Gemini')
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Save
output_path = r'c:/Users/700289/Downloads/Python/care-dx-app/docs/介護DXカカナイ_操作マニュアル.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
